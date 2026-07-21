"""
bridge/generate_proposal.py
实验方案 docx 生成器
========================
输入: 单体对 (醛 CAS, 胺 CAS) + 关联失败 ID
输出: experiment/proposals/{方案编号}_v{N}.docx

依赖: pip install python-docx
"""
import argparse
import csv
import datetime
import json
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---- 路径 ----
HERE = Path(__file__).parent.resolve()
PROJ = HERE.parent

REAGENT_DB = PROJ / 'experiment' / 'reagent_db.json'
FEEDBACK_DB = PROJ / 'experiment' / 'feedback_db.csv'
HISTORY = PROJ / 'experiment' / 'history' / 'index.json'
STRUCT_DIR = PROJ / 'experiment' / 'structure'
PROPOSALS_DIR = PROJ / 'experiment' / 'proposals'
PROPOSALS_DIR.mkdir(parents=True, exist_ok=True)

sys.path.insert(0, str(HERE))
from search_local_pdfs import search as rag_search, format_results_for_prompt
from query_graphrag import query as graphrag_query

# ---- 中文格式 ----
CN_FONT = '宋体'
EN_FONT = 'Times New Roman'
BODY_SIZE = 10.5  # 五号宋体


def _set_run_fonts(run, size=BODY_SIZE, bold=False):
    """设置 run 字体 + 中英文 + 字号 + 加粗"""
    run.font.size = Pt(size)
    run.font.name = EN_FONT
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), CN_FONT)
    rFonts.set(qn('w:ascii'), EN_FONT)
    rFonts.set(qn('w:hAnsi'), EN_FONT)
    if bold:
        run.bold = True


def setup_document_styles(doc):
    """全局样式设定:1.5 倍行距, 段前段后 0, 正文首行缩进 2 字符"""
    # Normal 样式
    style = doc.styles['Normal']
    style.font.size = Pt(BODY_SIZE)
    style.font.name = EN_FONT
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), CN_FONT)
    rFonts.set(qn('w:ascii'), EN_FONT)
    rFonts.set(qn('w:hAnsi'), EN_FONT)
    # 行距 1.5, 段前段后 0
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


def set_paragraph_format(p, first_line_indent=True):
    """对单个段落应用 1.5 倍行距, 段前段后 0"""
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    if first_line_indent:
        # 2 字符缩进 ≈ 0.74 cm (10.5pt × 2)
        pf.first_line_indent = Cm(0.74)
    else:
        pf.first_line_indent = Cm(0)


def set_cell_text(cell, text, bold=False, size=BODY_SIZE):
    cell.text = ''
    p = cell.paragraphs[0]
    set_paragraph_format(p, first_line_indent=False)  # 表格不缩进
    p.paragraph_format.line_spacing = 1.0  # 表格内紧凑
    run = p.add_run(text)
    _set_run_fonts(run, size=size, bold=bold)


def add_styled_paragraph(doc, text, size=BODY_SIZE, bold=False, alignment=None,
                       no_indent=False):
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line_indent=not no_indent)
    run = p.add_run(text)
    _set_run_fonts(run, size=size, bold=bold)
    if alignment is not None:
        p.alignment = alignment
    return p


def add_section_heading(doc, text, level=1):
    h = doc.add_heading('', level=level)
    set_paragraph_format(h, first_line_indent=False)
    # heading 自带 section spacing, 显式重设
    h.paragraph_format.space_before = Pt(6)
    h.paragraph_format.space_after = Pt(6)
    h.paragraph_format.line_spacing = 1.5
    run = h.add_run(text)
    _set_run_fonts(run, size=14 - level*2, bold=True)
    return h


# ---- 数据访问 ----
def load_reagent_db():
    with open(REAGENT_DB, encoding='utf-8') as f:
        return json.load(f)


def find_reagent(reagents, cas):
    for r in reagents['reagents']:
        if r['cas'] == cas:
            return r
    return None


def format_graphrag_results(qr):
    """把 GraphRAG 输出转为可插入 docx 的中文文本"""
    lines = []
    lines.append(f'查询: {qr["query"]}')
    lines.append(f'关键词: {qr["keywords"]}')
    lines.append(f'命中: {qr["summary"]["n_reactions"]} 反应, {qr["summary"]["n_literatures"]} 文献')
    lines.append('')
    lines.append('TOP 5 相关反应:')
    for i, h in enumerate(qr['reactions'][:5], 1):
        r = h['data']
        ald = r.get('aldehyde_name', '?')[:50]
        ami = r.get('amine_name', '?')[:50]
        solv = r.get('solvent', '?')[:60]
        temp = r.get('temperature', '?')
        outcome = r.get('outcome', '?')
        score = h['score']
        lines.append(f'  {i}. [{score}★] {ald} + {ami}')
        lines.append(f'     溶剂: {solv} | 温度: {temp} | 产物: {outcome}')
    lines.append('')
    lines.append('TOP 5 相关文献:')
    for i, h in enumerate(qr['literatures'][:5], 1):
        l = h['data']
        journal = l.get('journal', '?')
        system = l.get('system', '?')[:80]
        innovation = l.get('innovation', '?')[:150]
        score = h['score']
        lines.append(f'  {i}. [{score}★] [{journal}]')
        lines.append(f'     体系: {system}')
        lines.append(f'     创新: {innovation}')
    return '\n'.join(lines)


def load_feedback_db():
    with open(FEEDBACK_DB, encoding='utf-8-sig', newline='') as f:
        return list(csv.DictReader(f))


def load_history():
    with open(HISTORY, encoding='utf-8') as f:
        return json.load(f)


# ---- 文档生成 ----
def generate(aldehyde_cas, amine_cas, related_failure_id=None,
             version=2, target_node_prefix=None):
    """生成实验方案 docx

    aldehyde_cas, amine_cas: 单体 CAS
    related_failure_id: 关联的失败 feedback 编号 (如 'COF-TAPT-2026-06-24-C-v1')
    version: v{N}
    target_node_prefix: 'TAPT' 或 'TAPB' (中心核前缀)
    """
    reagents = load_reagent_db()
    fb_db = load_feedback_db()
    history = load_history()

    ald = find_reagent(reagents, aldehyde_cas)
    ami = find_reagent(reagents, amine_cas)
    if not ald or not ami:
        raise ValueError(f'CAS 未在试剂库中找到: {aldehyde_cas} 或 {amine_cas}')

    # 节点前缀推断
    if target_node_prefix is None:
        if ami['cas'] == '14544-47-9':
            target_node_prefix = 'TAPT'
        elif ami['cas'] == '118727-34-7':
            target_node_prefix = 'TAPB'
        elif ami['cas'] == '341-58-2' or ald['cas'] == '443922-06-3':
            target_node_prefix = 'TFPT'
        else:
            target_node_prefix = 'NODE'

    # 方案编号
    today = datetime.date.today().isoformat()
    prop_id = f'COF-{target_node_prefix}-{today}-{ald["cas"]}_{ami["cas"]}-v{version}'

    # 关联失败
    related_failure = None
    if related_failure_id:
        for r in fb_db:
            if r.get('方案编号') == related_failure_id:
                related_failure = r
                break

    # RAG 检索 (核心知识库 + tianxuan 全库)
    rag_query = f'{target_node_prefix} {ami.get("name_short", "")} 亚胺COF 膜合成 反应条件'
    rag_results = rag_search({
        'aldehyde_cas': aldehyde_cas,
        'amine_cas': amine_cas,
        'keywords': [target_node_prefix, 'CF3', '膜'],
        'query_text': rag_query,
        'use_tianxuan': True,
        'top_k_tianxuan': 8,
        'tianxuan_min_sim': 0.65,
    })
    rag_text = format_results_for_prompt(rag_results)

    # GraphRAG 检索 (补充 954 篇文献的实体关联)
    graphrag_query_text = f'{target_node_prefix} {ami["name_short"]} 膜 120°C'
    graphrag_results = graphrag_query(graphrag_query_text, verbose=False)
    graphrag_text = format_graphrag_results(graphrag_results)

    # ---- 创建文档 ----
    doc = Document()
    setup_document_styles(doc)

    # ===== 标题 =====
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(title_p, first_line_indent=False)
    title_run = title_p.add_run(f'实验方案_v{version}:{ald["name_short"] or ald["cas"]} + {ami["name_short"] or ami["cas"]}')
    _set_run_fonts(title_run, size=16, bold=True)

    # 标题块均不缩进
    add_styled_paragraph(doc, f'方案编号:{prop_id}', size=BODY_SIZE, alignment=WD_ALIGN_PARAGRAPH.CENTER, no_indent=True)
    add_styled_paragraph(doc, f'生成日期:{today}', size=BODY_SIZE, alignment=WD_ALIGN_PARAGRAPH.CENTER, no_indent=True)
    add_styled_paragraph(doc, f'迭代依据:{related_failure_id or "(无关联失败)"}', size=BODY_SIZE, alignment=WD_ALIGN_PARAGRAPH.CENTER, no_indent=True)
    doc.add_paragraph()  # 空行

    # ===== 0. 元信息表 =====
    add_section_heading(doc, '0. 元信息', level=1)
    meta_table = doc.add_table(rows=0, cols=2)
    meta_table.style = 'Light Grid Accent 1'
    meta_rows = [
        ('方案编号', prop_id),
        ('版本', f'v{version}'),
        ('生成日期', today),
        ('迭代依据', related_failure_id or '无'),
        ('节点前缀', target_node_prefix),
        ('醛单体', f'{ald["name_short"] or ald["cas"]} (CAS {ald["cas"]})'),
        ('胺单体', f'{ami["name_short"] or ami["cas"]} (CAS {ami["cas"]})'),
        ('试剂状态 (醛)', ald.get('status', '未知')),
        ('试剂状态 (胺)', ami.get('status', '未知')),
    ]
    for k, v in meta_rows:
        row = meta_table.add_row().cells
        set_cell_text(row[0], k, bold=True)
        set_cell_text(row[1], v)

    # ===== 1. 单体结构 =====
    add_section_heading(doc, '1. 单体结构', level=1)
    add_styled_paragraph(doc, f'醛单体 {ald["name_short"] or ald["cas"]}:')
    ald_img = STRUCT_DIR / f'{ald["cas"]}.png'
    if ald_img.exists():
        doc.add_picture(str(ald_img), width=Inches(2.5))
        last_p = doc.paragraphs[-1]
        last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_styled_paragraph(doc, f'CAS: {ald["cas"]}    MW: {ald.get("mw","?")}')
    doc.add_paragraph()

    add_styled_paragraph(doc, f'胺单体 {ami["name_short"] or ami["cas"]}:')
    ami_img = STRUCT_DIR / f'{ami["cas"]}.png'
    if ami_img.exists():
        doc.add_picture(str(ami_img), width=Inches(2.5))
        last_p = doc.paragraphs[-1]
        last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_styled_paragraph(doc, f'CAS: {ami["cas"]}    MW: {ami.get("mw","?")}')

    # ===== 2. 研究背景与科学问题 =====
    add_section_heading(doc, '2. 研究背景与科学问题', level=1)
    if related_failure:
        bg_text = (
            f'本方案基于历史失败 {related_failure_id} 的迭代改进。\n\n'
            f'原失败:{related_failure.get("失败现象描述","?")[:200]}\n\n'
            f'原 Class:{related_failure.get("失败Class","?")}\n'
            f'原类型:{related_failure.get("根因Type","?")}\n\n'
            f'本方案要回答的科学问题:'
        )
        if target_node_prefix == 'TAPT':
            bg_text += f'\n- TAPT 中心核(vs 苯核)是否能在该 {ald["name_short"]} + {ami["name_short"]} 体系下促进连续膜形成?'
        elif target_node_prefix == 'TFPT':
            bg_text += f'\n- TFPT 三嗪醛 + 该胺体系的成膜条件窗口是什么?'
        add_styled_paragraph(doc, bg_text)
    else:
        add_styled_paragraph(doc, '本方案为正向基线方案。背景与科学问题待用户补充。')

    # ===== 3. 关联历史失败 =====
    if rag_results.get('feedback_matches') or related_failure:
        add_section_heading(doc, '3. 关联历史失败', level=1)

        if related_failure:
            add_styled_paragraph(doc, '▌本次迭代的直接相似案例', bold=True, size=11)
            rf_table = doc.add_table(rows=0, cols=2)
            rf_table.style = 'Light Grid Accent 1'
            rf_rows = [
                ('失败 ID', related_failure_id),
                ('醛 CAS', related_failure.get('醛CAS','?')),
                ('胺 CAS', related_failure.get('胺CAS','?')),
                ('失败 Class', related_failure.get('失败Class','?')),
                ('根因 Type', related_failure.get('根因Type','?')),
                ('置信度', '高(直接相似)' if target_node_prefix in related_failure.get('关联历史失败','') else '低(外推)'),
                ('处置', related_failure.get('下一轮建议','?')),
                ('引用', '[本地: experiment/feedback_db.csv, 行 ' + related_failure_id + ']'),
            ]
            for k, v in rf_rows:
                row = rf_table.add_row().cells
                set_cell_text(row[0], k, bold=True)
                set_cell_text(row[1], v)

        fb_matches = rag_results.get('feedback_matches', [])
        if fb_matches and len(fb_matches) > 1:
            add_styled_paragraph(doc, '', size=4)
            add_styled_paragraph(doc, '▌RAG 检索到的其他相关失败', bold=True, size=11)
            for r in fb_matches:
                if r.get('方案编号') == related_failure_id:
                    continue
                add_styled_paragraph(
                    doc,
                    f'• {r.get("方案编号")} | 醛+胺: {r.get("醛CAS","?")}+{r.get("胺CAS","?")} | '
                    f'Class {r.get("失败Class","?")} | 现象: {r.get("失败现象描述","?")[:80]}',
                    size=9)

    # ===== 4. 实验设计与分组 =====
    add_section_heading(doc, '4. 实验设计与分组', level=1)
    add_styled_paragraph(doc, '本方案建议做以下对照实验组(每个实验组对应一个科学问题):')
    design_table = doc.add_table(rows=1, cols=5)
    design_table.style = 'Light Grid Accent 1'
    hdr = design_table.rows[0].cells
    headers = ['实验编号', '固定条件', '变量', '科学问题', '阳性对照']
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True)

    design_rows = [
        (
            'G1-S1',
            f'醛 = {ald["name_short"]}',
            f'胺 = {ami["name_short"]}',
            f'本体系基线可行性',
            'AMCOF-1: TFPT + BD-CF3'
        ),
        (
            'G1-S2',
            f'醛 = {ald["name_short"]}',
            f'胺 = {ami["name_short"]}',
            '延长反应时间 72h → 96h',
            'G1-S1'
        ),
        (
            'G1-S3',
            f'醛 = {ald["name_short"]}',
            f'胺 = {ami["name_short"]}',
            '苯胺 8.2 μL → 13.7 μL',
            'G1-S1'
        ),
    ]
    for r in design_rows:
        cells = design_table.add_row().cells
        for i, v in enumerate(r):
            set_cell_text(cells[i], v)

    # ===== 5. 通用实验流程 =====
    add_section_heading(doc, '5. 通用实验流程 (10 步模板)', level=1)
    add_styled_paragraph(doc, '以下流程可复用,仅各组具体参数差异化。', size=BODY_SIZE, no_indent=True)
    alde_label = ald['name_short'] or ald['cas']
    amine_label = ami['name_short'] or ami['cas']
    steps = [
        ('步骤 1', 'Pyrex 管准备', '35 mL Pyrex 玻璃管依次用丙酮、乙醇、去离子水超声清洗各 10 min,120°C 干燥备用。'),
        ('步骤 2', '投料:醛基单体 + 苯胺',
            f'向 Pyrex 管中依次加入醛基单体({alde_label})、溶剂甲苯 + 氯仿、苯胺(按醛基当量 1:1)。'
            '【单次连续加原则】苯胺与醛基同步溶解,立即进入步骤 3。'),
        ('步骤 3', '立即加入胺基单体(不间隔)',
            f'将含胺基的单体({amine_label})溶解于氯仿中,超声 5 min 后立即加入步骤 2 的混合液中。'
            '【关键】加入胺单体后不超过 5 min 必须进入步骤 4 加酸--避免长时间停留导致苯胺-醛亚胺中间体被胺交换取代不完全。'),
        ('步骤 4', '最后加入乙酸催化剂',
            '沿管壁缓慢加入 6.0 M 乙酸水溶液(0.20 mL,单点一次性加入)。'
            '【禁 序】乙酸必须最后加,先加乙酸会引发醛胺直接不可逆反应破坏可控性。'),
        ('步骤 5', '超声混合', '混合物超声 10-15 min 至完全溶解或均匀分散。'),
        ('步骤 6', '密封与加热', 'PTFE 内衬螺旋盖密封,置于预设温度的油浴中静置反应。'),
        ('步骤 7', '反应监测', '定期观察玻璃管内壁液面以上区域是否有薄膜形成,记录膜颜色和形态变化。'),
        ('步骤 8', '膜收集', '反应结束后冷却至室温,用金属刮刀小心将玻璃壁上的薄膜剥离。'),
        ('步骤 9', '洗涤', '依次用丙酮和四氢呋喃(THF)洗涤膜样品各 3 次,去除未反应单体和低聚物。'),
        ('步骤 10', '干燥与称量', '室温自然干燥(或 60°C 真空干燥过夜),称量膜质量并计算分离产率。'),
    ]
    for step_no, name, detail in steps:
        p = doc.add_paragraph()
        run_no = p.add_run(f'{step_no}  ')
        _set_run_fonts(run_no, size=10, bold=True)
        run_name = p.add_run(f'{name}\n')
        _set_run_fonts(run_name, size=10, bold=True)
        run_detail = p.add_run(f'  {detail}')
        _set_run_fonts(run_detail, size=9)

    # ===== 6. 各组详细方案 =====
    add_section_heading(doc, '6. 各组详细方案', level=1)
    group_table = doc.add_table(rows=1, cols=5)
    group_table.style = 'Light Grid Accent 1'
    ghdr = group_table.rows[0].cells
    gheaders = ['编号', '单体组合', '温度/时间', '溶剂体系', '苯胺 / 乙酸']
    for i, h in enumerate(gheaders):
        set_cell_text(ghdr[i], h, bold=True)

    groups = [
        ('G1-S1', f'{ald["name_short"]} + {ami["name_short"]}',
            '120°C / 72h',
            '甲苯 0.6 mL / 氯仿 0.4 mL',
            '苯胺 13.7 μL / 6.0M HOAc 0.20 mL'),
        ('G1-S2', f'{ald["name_short"]} + {ami["name_short"]}',
            '120°C / 96h (延长)',
            '甲苯 0.6 mL / 氯仿 0.4 mL',
            '苯胺 13.7 μL / 6.0M HOAc 0.20 mL'),
        ('G1-S3', f'{ald["name_short"]} + {ami["name_short"]}',
            '120°C / 72h',
            '甲苯 0.6 mL / 氯仿 0.4 mL',
            '苯胺 13.7 μL (提高) / 6.0M HOAc 0.20 mL'),
    ]
    for g in groups:
        cells = group_table.add_row().cells
        for i, v in enumerate(g):
            set_cell_text(cells[i], v, bold=(i==0))

    # ===== 7. 关键参数说明 =====
    add_section_heading(doc, '7. 关键参数说明', level=1)
    params = [
        ('溶剂选择原则',
            '基于侯老师机理,溶剂需 (I) 低沸点快速蒸发形成超饱和液层;(II) 良好玻璃润湿性;(III) 对单体的充分溶解能力。'
            '甲苯(bp 110.6°C)提供润湿性;氯仿(bp 61.2°C)提供高挥发性。'),
        ('苯胺调制量',
            f'苯胺通过可逆反应封端醛基活性位点,增强可逆性促进"错误校验"和自我修复。'
            f'苯胺:-CHO = 1:1(本体系对 TFPT 三醛节点即 3 eq)。'
            f'若反应活性偏低,优先通过延长时间/调整溶剂/催化剂解决,而非增加苯胺--4-5 eq 时 fwhm 增大、膜质量下降。'),
        ('乙酸催化剂',
            '6.0 M 乙酸水溶液提供适宜 pH 催化席夫碱缩合。乙酸 pKa 与苯胺匹配,既能催化又能维持可逆性。'),
    ]
    for k, v in params:
        add_styled_paragraph(doc, f'{k}', bold=True, size=10)
        add_styled_paragraph(doc, v, size=9)

    # ===== 8. 失败排查与修正 =====
    add_section_heading(doc, '8. 失败排查与修正', level=1)
    if related_failure:
        add_styled_paragraph(
            doc,
            f'针对历史失败 {related_failure_id} (Class {related_failure.get("失败Class","?")}) 的根因,本次方案预设了对应修正:',
            size=10
        )
        for rec in (related_failure.get('下一轮建议','') or '').split('\n'):
            if rec.strip():
                add_styled_paragraph(doc, f'  • {rec.strip()}', size=9)
    else:
        add_styled_paragraph(doc, '常见问题排查参考 v3.9 §4.3。', size=10)

    add_styled_paragraph(doc, '', size=4)
    add_styled_paragraph(doc, '▌本方案预设的失败应对(来自 failure_playbook)', bold=True, size=11)
    playbooks = [
        ('A1 类(自反应法)失败模式 α', '若 72h 膜不连续 → 缩短自反应时间 12h → 6h;或加醛前降温'),
        ('A 类（操作错）失败模式', '严格按加料顺序：醛+苯胺 → 立即胺 → 最后乙酸；乙酸用 6.0M 而非 18M；苯胺按醛当量'),
        ('C 类(粉红色油状物)', '加酸分两次(0.3 + 0.3 mL 间隔 1h);或将乙酸降到 3.0 M'),
    ]
    for name, action in playbooks:
        add_styled_paragraph(doc, f'• {name}', bold=True, size=9)
        add_styled_paragraph(doc, f'  {action}', size=9)

    # ===== 9. 表征方案 =====
    add_section_heading(doc, '9. 表征方案', level=1)
    char_table = doc.add_table(rows=1, cols=3)
    char_table.style = 'Light Grid Accent 1'
    chdr = char_table.rows[0].cells
    for i, h in enumerate(['表征', '目的', '关键判定']):
        set_cell_text(chdr[i], h, bold=True)
    char_rows = [
        ('PXRD', '评估膜结晶度(fwhm100)', '对照 AMCOF-1, 目标 fwhm ≤ 0.16°'),
        ('SEM', '膜表面形貌与连续性', '气相侧/玻璃侧差异;连续覆盖'),
        ('AFM', '膜厚测量多点取样', '涂覆区域是否均匀'),
        ('FT-IR', '亚胺键 C=N ~1631 cm-1', '确认醛基消耗完全度'),
        ('BET (N2, 77K)', '比表面积和孔径分布', '> 800 m2/g 为佳'),
    ]
    for r in char_rows:
        cells = char_table.add_row().cells
        for i, v in enumerate(r):
            set_cell_text(cells[i], v)

    # ===== 10. RAG 检索结果 =====
    add_section_heading(doc, '10. RAG 检索结果(生成时刻快照)', level=1)
    add_styled_paragraph(doc, rag_text or '(无匹配)', size=8)

    # ===== 10b. GraphRAG 实体检索结果 =====
    add_section_heading(doc, '10b. GraphRAG 实体检索 (954 篇文献交叉)', level=1)
    add_styled_paragraph(doc, graphrag_text or '(无匹配)', size=8)

    # ===== 11. 参考文献 =====
    add_section_heading(doc, '11. 参考文献', level=1)
    refs = [
        '侯盛怀等, Angew. Chem. Int. Ed. 2025, 64, e202421555 (扩散/调制剂双介导固-液/气三相界面合成)',
        '[本地: 实验/文章/侯老师实验/侯盛怀德国应化.pdf]',
    ]
    if related_failure_id:
        refs.append(f'[本地: experiment/feedback_db.csv, 行 {related_failure_id}]')
    refs.append('本方案基于 v3.9 实验方案_含氟COF薄膜合成_v3.9_20260626.docx 模板')

    for ref in refs:
        add_styled_paragraph(doc, f'  • {ref}', size=9)

    # ===== 保存 =====
    out_path = PROPOSALS_DIR / f'{prop_id}.docx'
    doc.save(out_path)
    print(f'✓ 已生成: {out_path}')
    return out_path


if __name__ == '__main__':
    parser = argparse.ArgumentParser()
    parser.add_argument('--ald', required=True, help='醛 CAS')
    parser.add_argument('--amine', required=True, help='胺 CAS')
    parser.add_argument('--related', default=None, help='关联失败 ID')
    parser.add_argument('--version', type=int, default=2)
    parser.add_argument('--node', default=None)
    args = parser.parse_args()

    generate(args.ald, args.amine, args.related, args.version, args.node)
