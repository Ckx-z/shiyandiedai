"""
bridge/index_knowledge.py
==========================
知识库 PDF/docx 解析 + MiniMax embedding + 索引

用法:
  # 索引核心 10 篇
  python index_knowledge.py --core
  
  # 索引所有 PDF/docx (慢)
  python index_knowledge.py --full
  
  # 索引特定目录
  python index_knowledge.py --dir "知识库/侯老师实验"
  
  # 查询
  python index_knowledge.py --query "TFPT 三嗪醛 成膜条件"
"""
import os
import sys
import json
import time
import argparse
import requests
import hashlib
from pathlib import Path
import re

# ---- 路径 ----
PROJ = Path(__file__).resolve().parent.parent
KNOWLEDGE_ROOT = PROJ / '知识库'
INDEX_FILE = PROJ / 'bridge' / 'knowledge_index.jsonl'
META_FILE = PROJ / 'bridge' / 'knowledge_meta.json'

# ---- MiniMax embedding API ----
EMBEDDING_API = 'https://api.minimax.chat/v1/embeddings'
EMBEDDING_MODEL = 'embo-01'
EMBEDDING_DIM = 1536
CHUNK_SIZE = 500
CHUNK_OVERLAP = 100


def get_api_key():
    key = os.environ.get('MINIMAX_API_KEY')
    if not key:
        raise RuntimeError('MINIMAX_API_KEY 未设置 (运行 _tmp/set_api_env.ps1)')
    return key


def compute_embedding(texts, embed_type='db'):
    """调用 MiniMax embedding API"""
    api_key = get_api_key()
    r = requests.post(
        EMBEDDING_API,
        headers={'Authorization': f'Bearer {api_key}', 'Content-Type': 'application/json'},
        json={'model': EMBEDDING_MODEL, 'texts': texts, 'type': embed_type},
        timeout=60,
    )
    r.raise_for_status()
    j = r.json()
    if j.get('base_resp', {}).get('status_code', 0) != 0:
        raise RuntimeError(f'Embedding API error: {j.get("base_resp", {})}')
    return j['vectors']


def parse_pdf(path):
    """pdfplumber 抽 PDF 文本"""
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                t = page.extract_text() or ''
                text_parts.append(t)
        return '\n\n'.join(text_parts)
    except Exception as e:
        print(f'  PDF 解析失败: {e}')
        return ''


def parse_docx(path):
    """python-docx 抽 docx 文本"""
    try:
        from docx import Document
        doc = Document(str(path))
        text_parts = [p.text for p in doc.paragraphs if p.text.strip()]
        # 表格
        for tbl in doc.tables:
            for row in tbl.rows:
                cells = [c.text.strip() for c in row.cells]
                text_parts.append(' | '.join(cells))
        return '\n\n'.join(text_parts)
    except Exception as e:
        print(f'  DOCX 解析失败: {e}')
        return ''


def chunk_text(text, chunk_size=CHUNK_SIZE, overlap=CHUNK_OVERLAP):
    """滑动窗口分块"""
    text = text.strip()
    if not text:
        return []
    # 按段落优先
    paragraphs = re.split(r'\n\s*\n', text)
    chunks = []
    current = ''
    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
        # 单段超长则切
        if len(para) > chunk_size:
            if current:
                chunks.append(current)
                current = ''
            for i in range(0, len(para), chunk_size - overlap):
                chunks.append(para[i:i + chunk_size])
            continue
        # 累积到 chunk_size
        if len(current) + len(para) + 2 > chunk_size:
            chunks.append(current)
            current = para
        else:
            current = (current + '\n\n' + para) if current else para
    if current:
        chunks.append(current)
    return chunks


def index_files(file_paths, batch_size=8):
    """批量索引一组文件"""
    all_records = []
    total_chunks = 0

    for path in file_paths:
        path = Path(path)
        if not path.exists():
            print(f'  ! 不存在: {path}')
            continue

        print(f'索引: {path.name}')

        if path.suffix.lower() == '.pdf':
            text = parse_pdf(path)
        elif path.suffix.lower() == '.docx':
            text = parse_docx(path)
        else:
            continue

        if not text.strip():
            print(f'  ! 空文本, 跳过')
            continue

        chunks = chunk_text(text)
        total_chunks += len(chunks)
        print(f'  分块: {len(chunks)} 块')

        # 批量 embedding (每批 batch_size)
        for i in range(0, len(chunks), batch_size):
            batch = chunks[i:i + batch_size]
            try:
                vectors = compute_embedding(batch, embed_type='db')
            except Exception as e:
                print(f'  Embedding 失败: {e}')
                continue

            for j, vec in enumerate(vectors):
                chunk_idx = i + j
                chunk_text_clean = batch[j].replace('\n', ' ').strip()
                record = {
                    'path': str(path.relative_to(PROJ)) if path.is_relative_to(PROJ) else str(path),
                    'chunk_id': f'{path.stem}_{chunk_idx:03d}',
                    'text': chunk_text_clean[:600],  # 截断存入 (向量才是主数据)
                    'vector': vec,
                    'dim': len(vec),
                }
                all_records.append(record)

    print(f'\n总块数: {total_chunks}, 总记录: {len(all_records)}')
    return all_records


def save_index(records):
    with open(INDEX_FILE, 'w', encoding='utf-8') as f:
        for r in records:
            f.write(json.dumps(r, ensure_ascii=False) + '\n')
    print(f'索引写入: {INDEX_FILE} ({len(records)} 条)')

    # 写 meta
    meta = {
        'indexed_at': time.strftime('%Y-%m-%d %H:%M:%S'),
        'total_records': len(records),
        'embedding_model': EMBEDDING_MODEL,
        'dim': EMBEDDING_DIM,
        'chunk_size': CHUNK_SIZE,
        'chunk_overlap': CHUNK_OVERLAP,
        'files': sorted(set(r['path'] for r in records)),
    }
    with open(META_FILE, 'w', encoding='utf-8') as f:
        json.dump(meta, f, ensure_ascii=False, indent=2)
    print(f'元数据: {META_FILE}')


def search_index(query_text, top_k=10, min_similarity=0.5):
    """检索索引

    1. 计算 query embedding (type=query)
    2. 加载所有索引
    3. 算 cosine similarity (纯 Python, 不依赖 numpy)
    4. 返回 Top-K
    """
    import math

    if not INDEX_FILE.exists():
        print('索引不存在, 请先运行 --core / --full')
        return []

    print(f'查询: {query_text}')
    query_vec = compute_embedding([query_text], embed_type='query')[0]

    records = []
    with open(INDEX_FILE, encoding='utf-8') as f:
        for line in f:
            records.append(json.loads(line))

    print(f'索引 {len(records)} 条')

    # 预计算 query norm
    q_norm = math.sqrt(sum(x * x for x in query_vec)) + 1e-10

    sims = []
    for r in records:
        v = r['vector']
        # 算 dot product 和 v norm
        dot = sum(a * b for a, b in zip(query_vec, v))
        v_norm = math.sqrt(sum(x * x for x in v)) + 1e-10
        sim = dot / (q_norm * v_norm)
        if sim >= min_similarity:
            sims.append((sim, r))
    sims.sort(reverse=True)
    return sims[:top_k]


def format_results(results, max_items=10):
    lines = []
    for sim, r in results:
        lines.append(f'### 相似度 {sim:.3f}')
        lines.append(f'  来源: {r["path"]} (chunk {r["chunk_id"]})')
        lines.append(f'  文本: {r["text"][:300]}')
        lines.append('')
    return '\n'.join(lines)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument('--core', action='store_true', help='索引核心 10 篇')
    parser.add_argument('--full', action='store_true', help='索引全量 (慢)')
    parser.add_argument('--dir', help='索引特定目录')
    parser.add_argument('--files', nargs='+', help='索引特定文件')
    parser.add_argument('--query', help='查询问题')
    parser.add_argument('--top-k', type=int, default=10)
    parser.add_argument('--min-sim', type=float, default=0.5)
    args = parser.parse_args()

    # 查询模式
    if args.query:
        results = search_index(args.query, top_k=args.top_k, min_similarity=args.min_sim)
        print(format_results(results))
        return

    if args.core:
        files = [
            KNOWLEDGE_ROOT / '侯老师实验' / '侯盛怀德国应化.pdf',
            KNOWLEDGE_ROOT / '侯老师实验' / 'sl-884-anie202421555-sup-0001-misc_information.pdf',
            KNOWLEDGE_ROOT / '机器学习' / 'Chemist-GuidedHuman−AIWorkflowforCovalentOrganic FrameworkSynthesis' / 'Chemist-GuidedHuman−AIWorkflowforCovalentOrganic FrameworkSynthesis.pdf',
            KNOWLEDGE_ROOT / '机器学习' / 'Chemist-GuidedHuman−AIWorkflowforCovalentOrganic FrameworkSynthesis' / 'ja5c20068_si_001 (1).pdf',
            KNOWLEDGE_ROOT / '机器学习' / 'nature' / 's41467-026-69549-z (1).pdf',
            KNOWLEDGE_ROOT / '机器学习' / 'nature' / '41467_2026_69549_MOESM1_ESM.pdf',
            KNOWLEDGE_ROOT / '文献阅读' / '调节剂与溶剂诱导聚合构筑功能化共价有机框架薄膜' / 's41467-024-55114-z.pdf',
            KNOWLEDGE_ROOT / '文献阅读' / '调节剂与溶剂诱导聚合构筑功能化共价有机框架薄膜' / '41467_2024_55114_MOESM2_ESM.pdf',
            KNOWLEDGE_ROOT / '自组装膜' / '自组装COF膜文献综述_v2.docx',
            KNOWLEDGE_ROOT / '机理' / 'Belowich-2012-Dynamic-imine-chemistry.pdf',
        ]
    elif args.dir:
        d = KNOWLEDGE_ROOT / args.dir
        files = list(d.rglob('*.pdf')) + list(d.rglob('*.docx'))
    elif args.files:
        files = [Path(f) for f in args.files]
    elif args.full:
        files = list(KNOWLEDGE_ROOT.rglob('*.pdf')) + list(KNOWLEDGE_ROOT.rglob('*.docx'))
        # 排除 _extracted/ 和 文本备份
        files = [f for f in files if '_extracted' not in str(f) and not str(f).endswith(('.txt', '.py', '.md'))]
    else:
        print('请指定 --core / --full / --dir / --files / --query')
        return

    print(f'将索引 {len(files)} 个文件:')
    for f in files[:5]:
        print(f'  - {f.name}')
    if len(files) > 5:
        print(f'  ... +{len(files) - 5} 个')
    print()

    records = index_files(files)
    save_index(records)


if __name__ == '__main__':
    main()