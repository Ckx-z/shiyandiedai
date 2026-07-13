"""
bridge/inspect_abcdef.py
========================
读取 实验ABCDEF.docx, 列出:
- 所有"实验X"标识
- 哪些有"结论:"标注 (已完成)
- 哪些没标注 (进行中)
- 哪些已入库 feedback_db.csv
- 哪些未入库 (新增完成, 提醒用户填 CSV)

不修改 ABCDEF.docx (只读巡查)
"""
import sys
import re
import csv
from pathlib import Path
from docx import Document

ABCDEF_PATH = Path(r'C:\Users\ckx\Desktop\实验\方案\实验ABCDEF.docx')
PROJ = Path(__file__).resolve().parent.parent
FEEDBACK_DB = PROJ / 'experiment' / 'feedback_db.csv'


def read_paragraphs():
    doc = Document(str(ABCDEF_PATH))
    return doc.paragraphs, doc.tables


def find_experiment_sections(paragraphs):
    """找出所有 '实验X' / '实验XN' 标识的位置 (如 A, A1, A8, D9)"""
    sections = []
    pattern = re.compile(r'实验([A-N]\d*)\b')
    for i, p in enumerate(paragraphs):
        txt = p.text.strip()
        if pattern.match(txt):
            m = pattern.match(txt)
            sections.append({'id': f'{m.group(1)}', 'start_idx': i, 'paragraphs': []})
    return sections


def find_section_conclusion(sections, paragraphs):
    """对每个 section，找最近的"结论:"段"""
    for s in sections:
        end_idx = next((sec['start_idx'] for sec in sections if sec['start_idx'] > s['start_idx']), len(paragraphs))
        for j in range(s['start_idx'], end_idx):
            if '结论' in paragraphs[j].text:
                # 抽取结论片段 (前 200 字)
                s['conclusion'] = paragraphs[j].text[:300].strip()
                s['completed'] = True
                break
        if 'conclusion' not in s:
            s['completed'] = False
            s['conclusion'] = ''
    return sections


def load_feedback_ids():
    if not FEEDBACK_DB.exists():
        return set()
    with open(FEEDBACK_DB, encoding='utf-8-sig', newline='') as f:
        rows = list(csv.DictReader(f))
    ids = set()
    for r in rows:
        prop_id = r.get('方案编号', '')
        # 提取 -A-, -B-, -A1-, -D9- 等
        for m in re.finditer(r'-([A-N]\d*)\b', prop_id):
            ids.add(m.group(1))
    return ids


def main():
    paragraphs, tables = read_paragraphs()
    print(f'=== ABCDEF.docx 巡查 ({ABCDEF_PATH.name}) ===')
    print(f'段落数: {len(paragraphs)}, 表格数: {len(tables)}')
    print()

    sections = find_experiment_sections(paragraphs)
    sections = find_section_conclusion(sections, paragraphs)

    feedback_ids = load_feedback_ids()

    completed = [s for s in sections if s['completed']]
    in_progress = [s for s in sections if not s['completed']]

    print(f'总实验: {len(sections)} 个')
    print(f'  ✓ 已完成 (有结论): {len(completed)} 个')
    print(f'  ⏳ 进行中 (无结论): {len(in_progress)} 个')
    print(f'  已入库 feedback_db.csv: {len(feedback_ids)} 个')
    print()

    print('=== 已完成 ===')
    for s in completed:
        in_db = '✓' if s['id'] in feedback_ids else '✗ 待入库'
        print(f'  实验{s["id"]} [{in_db}]')
        if s['conclusion']:
            print(f'    结论: {s["conclusion"][:150]}')
    print()

    if in_progress:
        print('=== 进行中 (无结论) ===')
        for s in in_progress:
            print(f'  实验{s["id"]}')

    # 找新增完成的
    new_completed = [s for s in completed if s['id'] not in feedback_ids]
    if new_completed:
        print()
        print(f'⚠ 发现 {len(new_completed)} 个新完成实验未入库:')
        for s in new_completed:
            print(f'  实验{s["id"]}: {s["conclusion"][:150]}')
        print()
        print('请在 experiment/feedback_db.csv 追加这些实验条目 (见 HOW_TO_FILL.md)')


if __name__ == '__main__':
    main()